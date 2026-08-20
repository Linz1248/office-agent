import docx
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

DOCX = r"e:\agent-project\office-agent\材料\项目文档_优化.docx"

doc = docx.Document(DOCX)

BODY_ASCII = "Times New Roman"
BODY_EAST = "宋体"
HEAD_EAST = "微软雅黑"


def set_style_fonts(style, ascii_font, east_font):
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_font)
    rfonts.set(qn("w:cs"), east_font)


# 1. Set fonts on styles
heading_names = {"Heading 1", "Heading 2", "Heading 3", "Title"}
for style in doc.styles:
    try:
        name = style.name
    except Exception:
        continue
    if name in heading_names:
        set_style_fonts(style, BODY_ASCII, HEAD_EAST)
    elif name in ("Normal", "Body Text", "First Paragraph", "Image Caption",
                  "Captioned Figure", "Compact", "Block Text", "Caption"):
        set_style_fonts(style, BODY_ASCII, BODY_EAST)

# 2. Center figure paragraphs + captions
for p in doc.paragraphs:
    if p.style.name in ("Captioned Figure", "Image Caption"):
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # center the table-caption paragraphs (text starts with '表 N｜')
    if p.text.strip().startswith("\u8868") and "\uff5c" in p.text:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 3. Style tables: grid borders + header shading
def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = borders.find(qn("w:" + edge))
        if el is None:
            el = OxmlElement("w:" + edge)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "BFBFBF")


for table in doc.tables:
    set_table_borders(table)
    # header row shading + bold
    for cell in table.rows[0].cells:
        set_cell_shading(cell, "DEEAF6")
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True

doc.save(DOCX)
print("polish done")