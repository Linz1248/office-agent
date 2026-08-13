"""document_extract 服务：文档全文提取与字段抽取。

- PDF：PaddleOCR(PPStructure) + contextgem(LLM) 抽取字段，带 OCR/LLM 缓存。
- Word(.docx)/Excel(.xlsx)：纯 Python(python-docx/openpyxl) 结构化解析为 Markdown。
- 旧版 .doc/.xls：LibreOffice headless 转为 OOXML 后再解析。
提供用户注册/登录、文件上传/删除、文档全文读取与字段抽取等接口。
"""
from __future__ import annotations

import os
import re
import json
import time
import uuid
import hashlib
import difflib
import sqlite3
import asyncio
import shutil
import tempfile
import traceback
from pathlib import Path
from typing import List, Any, Dict, Optional
from datetime import datetime, timedelta
from contextlib import asynccontextmanager

from fastapi import FastAPI, File, UploadFile, HTTPException, status, Depends
from fastapi.responses import JSONResponse, FileResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
from pydantic import BaseModel, EmailStr
from openpyxl import Workbook, load_workbook
from urllib.parse import quote
import bcrypt
from jose import JWTError, jwt
import fitz  # PyMuPDF
import uvicorn

from contextgem import (
    Document,
    DocumentLLM,
    JsonObjectConcept,
    JsonObjectExample,
    StringConcept,
)
from paddleocr import PPStructureV3
import paddle

import config
from config import (
    UPLOAD_DIR,
    DATA_DIR,
    OUTPUT_DIR,
    USERS_DB,
    CACHE_DB,
    CACHE_VERSION,
    SECRET_KEY,
    ALGORITHM,
    ACCESS_TOKEN_EXPIRE_MINUTES,
    DEVICE,
    SAT_MODEL_ID,
    OLLAMA_API_BASE,
    OLLAMA_MODEL,
    LLM_TIMEOUT,
    LLM_SEED,
    PP_STRUCTURE_MODEL_DIRS,
    PRETRAINED_MODELS_DIR,
    PORT,
    PUBLIC_BASE_URL,
)


# 密码加密规则：直接使用 bcrypt（passlib 1.7.4 与 bcrypt>=4 不兼容），$2b$ 哈希与旧 passlib 哈希互通
BCRYPT_ROUNDS = 12
security = HTTPBearer()

# OCR 引擎：从本地 pretrained_models 加载各子模块
paddle.set_device(DEVICE)
_pp_model_kwargs = {
    key: str(PRETRAINED_MODELS_DIR / subdir)
    for key, subdir in PP_STRUCTURE_MODEL_DIRS.items()
}
PPStructure = PPStructureV3(
    **_pp_model_kwargs,
    use_region_detection=True,
    use_table_recognition=True,
    use_doc_unwarping=True,
    use_chart_recognition=False,
    use_textline_orientation=False,
    use_seal_recognition=False,
    use_formula_recognition=False,
    use_doc_orientation_classify=False,
    text_det_thresh=0.5,
    device=DEVICE,
)

# LibreOffice 可用性：旧版 .doc/.xls 依赖 soffice 转换；缺失时仅旧版不可用，
# 现代格式 .docx/.xlsx 与 PDF 正常工作（启动期探测一次，避免每请求 which）。
LO_AVAILABLE: bool = shutil.which("soffice") is not None

# 允许 OCR 的图片扩展名
_IMAGE_EXTS = {".jpg", ".jpeg", ".png", ".bmp", ".webp"}

# 允许上传的文档扩展名（PDF + 现代/旧版 Office + 图片：图片经 /doc_upload 落盘后
# 可走 read_document / extract_document / extract_to_excel，与 PDF/Word/Excel 同构）
_ALLOWED_EXTS = {".pdf", ".docx", ".xlsx", ".doc", ".xls"} | _IMAGE_EXTS


# ------------------------------------------------------------------
# 缓存相关函数
# ------------------------------------------------------------------

def init_cache_db():
    """创建缓存表，启用 WAL 模式提升并发读性能。

    cache.db 位于 NFS 上，启用 WAL 偶发“database is locked”；通过
    get_cache_connection() 的 nolock=1 + busy_timeout 规避，WAL 设置失败时
    降级为默认日志模式，避免阻断服务启动。
    """
    with get_cache_connection() as conn:
        try:
            conn.execute("PRAGMA journal_mode=WAL")
        except sqlite3.OperationalError as e:
            print(f"[cache] 启用 WAL 失败，降级为默认日志模式：{e}")
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS ocr_cache (
                file_hash  TEXT,
                version    INTEGER,
                ocr_text   TEXT,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                PRIMARY KEY (file_hash, version)
            )
        """
        )
        conn.execute(
            """
            CREATE TABLE IF NOT EXISTS llm_cache (
                file_hash   TEXT,
                fields_hash TEXT,
                version     INTEGER,
                result      TEXT,
                created_at  TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
                PRIMARY KEY (file_hash, fields_hash, version)
            )
        """
        )
        conn.commit()


def compute_file_hash(file_path: str) -> str:
    """对文件二进制内容计算 SHA256，流式读取避免大文件占内存"""
    sha256 = hashlib.sha256()
    with open(file_path, "rb") as f:
        while chunk := f.read(8192):
            sha256.update(chunk)
    return sha256.hexdigest()


def compute_fields_hash(fields: list, enhance: bool,
                        fields_enhance: list, fields_template: dict) -> str:
    """将请求参数序列化后计算 SHA256，排序保证顺序无关"""
    data = {
        "fields": sorted(fields),
        "enhance": enhance,
        "fields_enhance": sorted(fields_enhance) if fields_enhance else [],
        "fields_template": fields_template if fields_template else {},
    }
    serialized = json.dumps(data, sort_keys=True, ensure_ascii=False)
    return hashlib.sha256(serialized.encode("utf-8")).hexdigest()


def get_ocr_cache(file_hash: str) -> Optional[str]:
    """查询文档文本缓存（ocr_cache 表为通用文本缓存，PDF/Word/Excel 全文共用，按 file_hash 命中），否则返回 None"""
    with get_cache_connection() as conn:
        row = conn.execute(
            "SELECT ocr_text FROM ocr_cache WHERE file_hash=? AND version=?",
            (file_hash, CACHE_VERSION),
        ).fetchone()
    return row[0] if row else None


def set_ocr_cache(file_hash: str, ocr_text: str):
    """写入文档文本缓存（PDF/Word/Excel 全文共用此缓存，按 file_hash 键）"""
    with get_cache_connection() as conn:
        conn.execute(
            "INSERT OR REPLACE INTO ocr_cache (file_hash, version, ocr_text) VALUES (?, ?, ?)",
            (file_hash, CACHE_VERSION, ocr_text),
        )
        conn.commit()


def get_llm_cache(file_hash: str, fields_hash: str) -> Optional[dict]:
    """查询 LLM 缓存，命中则返回提取结果字典，否则返回 None"""
    with get_cache_connection() as conn:
        row = conn.execute(
            "SELECT result FROM llm_cache WHERE file_hash=? AND fields_hash=? AND version=?",
            (file_hash, fields_hash, CACHE_VERSION),
        ).fetchone()
    return json.loads(row[0]) if row else None


def set_llm_cache(file_hash: str, fields_hash: str, result: dict):
    """写入 LLM 缓存"""
    with get_cache_connection() as conn:
        conn.execute(
            "INSERT OR REPLACE INTO llm_cache (file_hash, fields_hash, version, result) VALUES (?, ?, ?, ?)",
            (file_hash, fields_hash, CACHE_VERSION, json.dumps(result, ensure_ascii=False)),
        )
        conn.commit()


@asynccontextmanager
async def lifespan(app: FastAPI):
    # 启动阶段：初始化上传目录、数据目录与缓存数据库
    UPLOAD_DIR.mkdir(exist_ok=True)
    DATA_DIR.mkdir(exist_ok=True)
    OUTPUT_DIR.mkdir(exist_ok=True)
    init_cache_db()
    init_db()
    yield


app = FastAPI(title="文档抽取接口Demo", lifespan=lifespan)

# 跨域资源共享
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# 数据模型
class UserRegister(BaseModel):
    username: str
    email: EmailStr
    password: str
    confirmPassword: str
    admin_password: str


class UserLogin(BaseModel):
    username: str
    password: str


class User(BaseModel):
    id: int
    username: str
    email: str


class Token(BaseModel):
    access_token: str
    token_type: str
    expiresIn: int
    user: User


class DocEtractRequest(BaseModel):
    filename: str
    fields: List[str]
    enhance: bool
    fields_enhance: List[str]
    fields_template: dict


class DocTextRequest(BaseModel):
    filename: str


class PasswordChange(BaseModel):
    username: str
    old_password: str
    new_password: str
    confirm_password: str


# 数据库连接：nolock=1 绕过 NFS 文件锁不可用问题，busy_timeout 防止偶发冲突
def get_db_connection():
    conn = sqlite3.connect(f"file:{USERS_DB}?nolock=1", uri=True)
    conn.execute("PRAGMA busy_timeout=5000")
    return conn


def get_cache_connection():
    """cache.db 连接：同样使用 nolock=1 兼容 NFS，busy_timeout 规避偶发锁定冲突。"""
    conn = sqlite3.connect(f"file:{CACHE_DB}?nolock=1", uri=True)
    conn.execute("PRAGMA busy_timeout=5000")
    return conn


# 数据库初始化
def init_db():
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute(
            """
            CREATE TABLE IF NOT EXISTS users (
                id INTEGER PRIMARY KEY AUTOINCREMENT,
                username TEXT UNIQUE NOT NULL,
                hashed_password TEXT NOT NULL,
                email TEXT UNIQUE NOT NULL,
                created_at TIMESTAMP DEFAULT CURRENT_TIMESTAMP
            )
        """
        )
        conn.commit()


def _bcrypt_payload(password: str) -> bytes:
    """bcrypt 上限 72 字节，按官方建议截断；hash 与 verify 使用同一截断以保证一致。"""
    return password.encode("utf-8")[:72]


def verify_password(plain_password, hashed_password):
    return bcrypt.checkpw(_bcrypt_payload(plain_password), hashed_password.encode("utf-8"))


def get_password_hash(password):
    return bcrypt.hashpw(_bcrypt_payload(password), bcrypt.gensalt(rounds=BCRYPT_ROUNDS)).decode("utf-8")


# JWT相关函数
def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    if expires_delta:
        expire = datetime.utcnow() + expires_delta
    else:
        expire = datetime.utcnow() + timedelta(minutes=15)
    to_encode.update({"exp": expire, "sub": data["sub"]})  # 明确设置 sub 字段
    encoded_jwt = jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)
    return encoded_jwt


def verify_token(credentials: HTTPAuthorizationCredentials = Depends(security)):
    try:
        payload = jwt.decode(credentials.credentials, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise HTTPException(
                status_code=status.HTTP_401_UNAUTHORIZED,
                detail="登录信息已失效",
                headers={"WWW-Authenticate": "Bearer"},
            )
        return username
    except JWTError:
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="登录信息已失效",
            headers={"WWW-Authenticate": "Bearer"},
        )


# 数据库操作函数
def get_user_by_username(username: str):
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM users WHERE username=?", (username,))
        return cursor.fetchone()


def get_user_by_email(email: str):
    with get_db_connection() as conn:
        cursor = conn.cursor()
        cursor.execute("SELECT * FROM users WHERE email = ?", (email,))
        return cursor.fetchone()


def update_user(username: str, password_hash: str):
    with get_db_connection() as conn:
        cursor = conn.cursor()
        try:
            cursor.execute(
                "UPDATE users SET hashed_password = ? WHERE username = ?",
                (password_hash, username),
            )
            conn.commit()
            return cursor.rowcount > 0
        except Exception:
            return False


def create_user(username: str, email: str, password_hash: str):
    with get_db_connection() as conn:
        cursor = conn.cursor()
        try:
            cursor.execute(
                "INSERT INTO users (username, email, hashed_password) VALUES (?, ?, ?)",
                (username, email, password_hash),
            )
            conn.commit()
            return cursor.lastrowid
        except sqlite3.IntegrityError:
            return None


@app.get("/")
async def read_root(current_user: str = Depends(verify_token)):
    return {"message": "欢迎使用sv"}


# 注册接口
@app.post("/register", response_model=dict)
async def register(user: UserRegister):
    if get_user_by_username(user.username):
        raise HTTPException(status_code=400, detail="用户名已存在")
    if get_user_by_email(user.email):
        raise HTTPException(status_code=400, detail="邮箱已存在")
    password_hash = get_password_hash(user.password)
    admin = get_user_by_username("admin")
    if not admin or not verify_password(user.admin_password, admin[2]):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="管理员密码错误",
        )

    user_id = create_user(user.username, user.email, password_hash)
    if user_id is None:
        raise HTTPException(status_code=500, detail="创建用户失败")
    return {"message": "注册成功"}


# 登录接口
@app.post("/login", response_model=Token)
async def login(user: UserLogin):
    db_user = get_user_by_username(user.username)
    if not db_user or not verify_password(user.password, db_user[2]):
        raise HTTPException(
            status_code=status.HTTP_401_UNAUTHORIZED,
            detail="用户名或密码错误",
            headers={"WWW-Authenticate": "Bearer"},
        )
    access_token_expires = timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES)
    access_token = create_access_token(
        data={"sub": user.username}, expires_delta=access_token_expires
    )
    return {
        "access_token": access_token,
        "token_type": "bearer",
        "expiresIn": ACCESS_TOKEN_EXPIRE_MINUTES * 60000,
        "user": User(id=db_user[0], username=db_user[1], email=db_user[3]),
    }


# 密码修改接口
@app.post("/change_password", response_model=dict)
async def change_password(password_change: PasswordChange):
    user = get_user_by_username(password_change.username)
    if not user:
        raise HTTPException(status_code=400, detail="用户不存在")

    if not verify_password(password_change.old_password, user[2]):
        raise HTTPException(status_code=400, detail="原密码错误")
    new_password_hash = get_password_hash(password_change.new_password)
    update_res = update_user(user[1], new_password_hash)

    if not update_res:
        raise HTTPException(status_code=500, detail="密码修改失败")
    return {"message": "密码修改成功"}


# ------------------------------------------------------------------
# 单文件上传接口
# ------------------------------------------------------------------
@app.post("/doc_upload", summary="上传文档（PDF/Word/Excel）")
async def doc_upload(file: UploadFile = File(...), current_user: str = Depends(verify_token)):
    if not file.filename:
        raise HTTPException(status_code=400, detail="文件不能为空")

    suffix = Path(file.filename).suffix.lower()
    if suffix not in _ALLOWED_EXTS:
        raise HTTPException(
            status_code=400,
            detail="不支持的文件类型，支持 PDF/Word(.docx/.doc)/Excel(.xlsx/.xls)",
        )

    new_name = generate_filename(file.filename)  # <uuid4>_<timestamp>.<ext>
    dest_path = UPLOAD_DIR / new_name

    try:
        with dest_path.open("wb") as buffer:
            while chunk := await file.read(1024 * 1024):  # 每次 1 MB
                buffer.write(chunk)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"写入文件失败：{exc}")

    return JSONResponse(
        {
            "status_code": status.HTTP_200_OK,
            "message": "上传成功",
            "saved_name": new_name,
        }
    )


@app.delete("/doc_delete/{filename}", summary="删除指定文件")
async def doc_delete(filename: str, current_user: str = Depends(verify_token)):
    file_path = _resolve_upload(filename)
    if not os.path.isfile(file_path):
        raise HTTPException(status_code=400, detail="文件不存在")

    os.remove(file_path)

    return JSONResponse({"status_code": status.HTTP_200_OK, "message": "删除成功"})


# ------------------------------------------------------------------
# 工具函数：生成新文件名
# ------------------------------------------------------------------
def generate_filename(original_filename: str) -> str:
    """
    命名规则： <uuid4>_<timestamp>.<ext>
    例：9f2e1c4f-8b3a-4e6f-b2c1-0d9a5f7b1c3e_1719469876.pdf
    """
    ext = Path(original_filename).suffix  # 保留原始扩展名
    unique_str = uuid.uuid4().hex  # 32 位随机字符串
    ts = str(int(time.time()))  # 当前时间戳
    return f"{unique_str}_{ts}{ext}"


def _resolve_upload(filename: str, base: Path = UPLOAD_DIR) -> Path:
    """安全解析上传文件路径，统一防路径穿越（.. / 绝对路径 / 分隔符）。

    所有从请求取 filename 并拼接上传目录的端点（/doc_text、/doc_extract→do_extract、
    /doc_delete）都走此函数，避免每处重复实现路径校验。上传为扁平命名（uuid_时间戳.ext），
    故先拒任何路径分隔符，再做 resolve+is_relative_to 归属校验（双层防御）。
    返回 base 下解析后的绝对路径；filename 为空或越界时抛 HTTPException(400)。
    """
    name = (filename or "").strip()
    if not name:
        raise HTTPException(status_code=400, detail="文件名不能为空")
    if "/" in name or "\\" in name:
        raise HTTPException(status_code=400, detail="非法文件名")
    base_resolved = base.resolve()
    dest = (base / name).resolve()
    # 解析后必须仍位于上传目录内，防 .. 与绝对路径穿越
    if not dest.is_relative_to(base_resolved):
        raise HTTPException(status_code=400, detail="非法文件名")
    return dest


def extract_concepts_enhance(file: str, concepts: list, concepts_template: list):
    # 创建 Document 实例
    doc = Document(raw_text=file)
    doc.sat_model_id = SAT_MODEL_ID  # 从本地加载
    doc.paragraph_segmentation_mode = "sat"

    structure_dict = {field: str for field in concepts}

    doc_concept = JsonObjectConcept(
        name="自定义内容抽取",
        description="从文档中抽取用户指定的字段",
        structure=structure_dict,
        add_references=True,
        add_justifications=True,
        singular_occurrence=True,
        examples=concepts_template,
    )
    doc.add_concepts([doc_concept])

    llm = DocumentLLM(
        model=OLLAMA_MODEL,
        api_base=OLLAMA_API_BASE,
        output_language="adapt",
        timeout=LLM_TIMEOUT,
        seed=LLM_SEED,
    )

    extracted_concepts = llm.extract_concepts_from_document(doc)

    return extracted_concepts[0].extracted_items


def _clean_pdf(file_path: str) -> str:
    """返回清洁后的临时 PDF 路径，调用者负责删除"""
    doc_pdf = fitz.open(file_path)
    for page in doc_pdf:
        # 删除所有嵌入图片（水印、遮罩图层）
        for img in page.get_images(full=True):
            page.delete_image(img[0])
        # 删除所有注释/标注
        for annot in page.annots() or []:
            annot.delete()

    tmp = tempfile.NamedTemporaryFile(suffix=".pdf", delete=False)
    clean_path = tmp.name
    tmp.close()
    doc_pdf.save(clean_path, garbage=4, deflate=True)
    doc_pdf.close()
    return clean_path


def _extract_text_by_ppstructure(pdf_path: str) -> str:
    """用 PPStructure 提取 Markdown 并合并"""
    output = PPStructure.predict(input=pdf_path)
    markdown_list = [res.markdown for res in output]
    result = PPStructure.concatenate_markdown_pages(markdown_list)
    return result["markdown_texts"] if hasattr(result, "__getitem__") and "markdown_texts" in result else str(result)


def _merge_text(first: str, second: str) -> str:
    """
    把 second 中独有的段落插入到 first 的对应位置，保持顺序且不重复。
    返回合并后的整段文本。
    """
    # 1. 按空行拆成段落
    first_paras = [p.strip() for p in first.split("\n\n") if p.strip()]
    second_paras = [p.strip() for p in second.split("\n\n") if p.strip()]

    # 2. 用 difflib.SequenceMatcher 找到两段文本的段落对齐关系
    sm = difflib.SequenceMatcher(None, first_paras, second_paras, autojunk=False)
    merged = []
    i = j = 0  # i 扫描 first_paras，j 扫描 second_paras
    for tag, i1, i2, j1, j2 in sm.get_opcodes():
        if tag == "equal":
            merged.extend(first_paras[i1:i2])
            i, j = i2, j2
        elif tag == "delete":
            merged.extend(first_paras[i1:i2])
            i = i2
        elif tag == "insert":
            merged.extend(second_paras[j1:j2])
            j = j2
        elif tag == "replace":
            merged.extend(first_paras[i1:i2])
            for p in second_paras[j1:j2]:
                if not any(
                    difflib.SequenceMatcher(None, p, fp).ratio() > 0.85
                    for fp in first_paras[i1:i2]
                ):
                    merged.append(p)
            i, j = i2, j2

    # 3. 拼回字符串
    return "\n\n".join(merged)


def extract_pdf_text(file_path: str, file_hash: str | None = None) -> tuple:
    """
    OCR 提取 PDF 全文，复用 file_hash 缓存。

    与 extract_concepts 共享同一份 OCR 管线与缓存键，避免重复 OCR。
    供 /doc_text 接口与 extract_concepts 复用。

    Args:
        file_path: PDF 绝对路径。
        file_hash: 由调用方预先计算并传入，用于命中 OCR 缓存。

    Returns:
        tuple: (pdf_text, timing_info)
            timing_info 键: ocr_cache_hit / ocr_time / ocr_cache_query_time
    """
    pdf_text = None
    ocr_cache_hit = False
    ocr_time = 0.0
    ocr_query_time = 0.0

    if file_hash:
        t0 = time.perf_counter()
        cached_text = get_ocr_cache(file_hash)
        ocr_query_time = time.perf_counter() - t0
        if cached_text is not None:
            print(f"[OCR 缓存命中] file_hash={file_hash[:12]}... (查询耗时: {ocr_query_time*1000:.2f}ms)")
            pdf_text = cached_text
            ocr_cache_hit = True

    if not ocr_cache_hit:
        ocr_start = time.perf_counter()
        # ---------------- 第一次OCR：清洁 PDF ----------------
        clean_path = _clean_pdf(file_path)
        try:
            text_clean = _extract_text_by_ppstructure(clean_path)
        finally:
            try:
                os.remove(clean_path)
            except OSError:
                pass

        # ---------------- 第二次OCR：原始 PDF ----------------
        text_raw = _extract_text_by_ppstructure(file_path)

        # ---------------- 合并：仅保留第二次多出的文本 ----------------
        pdf_text = _merge_text(text_clean, text_raw)
        ocr_time = time.perf_counter() - ocr_start

        # ---------------- 写入 OCR 缓存 ----------------
        if file_hash:
            set_ocr_cache(file_hash, pdf_text)
            print(f"[OCR 缓存写入] file_hash={file_hash[:12]}... (OCR 耗时: {ocr_time:.4f}s)")

    timing_info = {
        "ocr_cache_hit": ocr_cache_hit,
        "ocr_time": round(ocr_time, 4),
        "ocr_cache_query_time": round(ocr_query_time * 1000, 2),
    }
    return pdf_text, timing_info


def extract_image_text(file_path: str, file_hash: str | None = None) -> tuple:
    """对图片做 OCR（PPStructure），复用 file_hash 文本缓存。

    PPStructureV3 直接接受图片路径（jpg/png/bmp/webp），无需清洁/双跑/合并。
    供 /image_text 接口复用，与 PDF/Word/Excel 全文共用 ocr_cache（按 file_hash 键）。
    """
    if file_hash:
        t0 = time.perf_counter()
        cached = get_ocr_cache(file_hash)
        query_time = time.perf_counter() - t0
        if cached is not None:
            print(f"[图片缓存命中] file_hash={file_hash[:12]}... (查询耗时: {query_time*1000:.2f}ms)")
            return cached, {"ocr_cache_hit": True, "ocr_time": 0.0,
                            "ocr_cache_query_time": round(query_time * 1000, 2)}
    else:
        query_time = 0.0

    t0 = time.perf_counter()
    text = _extract_text_by_ppstructure(file_path)
    ocr_time = time.perf_counter() - t0

    if file_hash:
        set_ocr_cache(file_hash, text)
        print(f"[图片缓存写入] file_hash={file_hash[:12]}... (OCR 耗时: {ocr_time:.4f}s)")

    return text, {
        "ocr_cache_hit": False,
        "ocr_time": round(ocr_time, 4),
        "ocr_cache_query_time": round(query_time * 1000, 2),
    }


# ------------------------------------------------------------------
# Word / Excel 结构化解析（纯 Python，复用 ocr_cache 文本缓存）
# timing 键与 extract_pdf_text 对齐（ocr_cache_hit / ocr_time /
# ocr_cache_query_time），便于 /doc_text 统一处理。
# ------------------------------------------------------------------

# 单元格文本上限，避免单个超长单元格撑爆输出
_CELL_TEXT_LIMIT = 2000
# 每个工作表最大行数，超出截断并标注（最终输出另受框架 tool_result_limit 兜底）
_SHEET_ROW_LIMIT = 1000


def _clip_cell(value: Any) -> str:
    """单元格值归一化为字符串，过长截断，换行替为空格，转义表格分隔符 | 。"""
    if value is None:
        return ""
    text = str(value)
    if len(text) > _CELL_TEXT_LIMIT:
        text = text[:_CELL_TEXT_LIMIT] + "…"
    return text.replace("\r", " ").replace("\n", " ").replace("|", "\\|")


def _table_to_markdown(rows: list[list[str]]) -> str:
    """将二维行数据渲染为 Markdown 表格。

    首行作表头，第二行 `---` 分隔，其余为数据行；空表返回空串。
    docx 表格与 xlsx 工作表共用此函数，避免重复渲染逻辑。
    """
    norm = [list(r) for r in rows]
    if not norm:
        return ""
    width = max(len(r) for r in norm)
    for r in norm:  # 补齐不等长行
        if len(r) < width:
            r.extend([""] * (width - len(r)))
    lines = ["| " + " | ".join(norm[0]) + " |",
             "|" + "|".join(["---"] * width) + "|"]
    for r in norm[1:]:
        lines.append("| " + " | ".join(r) + " |")
    return "\n".join(lines)


def extract_docx_text(file_path: str, file_hash: str | None = None) -> tuple:
    """读取 .docx 全文（python-docx），段落与表格转 Markdown。复用文本缓存。

    Args:
        file_path: .docx 绝对路径。
        file_hash: 由调用方预先计算并传入，用于命中文本缓存。

    Returns:
        tuple: (text, timing_info)
    """
    if file_hash:
        t0 = time.perf_counter()
        cached = get_ocr_cache(file_hash)
        query_time = time.perf_counter() - t0
        if cached is not None:
            print(f"[Word 缓存命中] file_hash={file_hash[:12]}... (查询耗时: {query_time*1000:.2f}ms)")
            return cached, {"ocr_cache_hit": True, "ocr_time": 0.0,
                            "ocr_cache_query_time": round(query_time * 1000, 2)}
    else:
        query_time = 0.0

    from docx import Document  # 局部导入：仅解析 Word 时加载
    t0 = time.perf_counter()
    doc = Document(file_path)
    parts: list[str] = []
    # doc.paragraphs 与 doc.tables 是两套平铺序列（表格不在 paragraphs 中），
    # 这里取「先所有段落、再所有表格」，对问答/摘要足够（表格内容不丢）。
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)
    for table in doc.tables:
        rows = [[_clip_cell(c.text) for c in row.cells] for row in table.rows]
        md = _table_to_markdown(rows)
        if md:
            parts.append(md)
    text = "\n\n".join(parts)
    parse_time = time.perf_counter() - t0

    if file_hash:
        set_ocr_cache(file_hash, text)
        print(f"[Word 缓存写入] file_hash={file_hash[:12]}... (解析耗时: {parse_time:.4f}s)")

    return text, {
        "ocr_cache_hit": False,
        "ocr_time": round(parse_time, 4),
        "ocr_cache_query_time": round(query_time * 1000, 2),
    }


def extract_xlsx_text(file_path: str, file_hash: str | None = None) -> tuple:
    """读取 .xlsx 全文（openpyxl），每个工作表渲染为 Markdown 表格。复用文本缓存。

    data_only=True 取公式缓存值；从未在 Excel/LibreOffice 打开过的公式单元格可能为空。
    每表最多 _SHEET_ROW_LIMIT 行、每单元格最多 _CELL_TEXT_LIMIT 字符，超出截断并标注。
    """
    if file_hash:
        t0 = time.perf_counter()
        cached = get_ocr_cache(file_hash)
        query_time = time.perf_counter() - t0
        if cached is not None:
            print(f"[Excel 缓存命中] file_hash={file_hash[:12]}... (查询耗时: {query_time*1000:.2f}ms)")
            return cached, {"ocr_cache_hit": True, "ocr_time": 0.0,
                            "ocr_cache_query_time": round(query_time * 1000, 2)}
    else:
        query_time = 0.0

    t0 = time.perf_counter()
    wb = load_workbook(file_path, read_only=True, data_only=True)
    parts: list[str] = []
    try:
        for ws in wb.worksheets:
            sheet_parts: list[str] = [f"## 工作表: {ws.title}"]
            rows: list[list[str]] = []
            truncated = False
            for i, row in enumerate(ws.iter_rows(values_only=True)):
                if i >= _SHEET_ROW_LIMIT:
                    truncated = True
                    break
                rows.append([_clip_cell(c) for c in row])
            md = _table_to_markdown(rows)
            if md:
                sheet_parts.append(md)
            else:
                sheet_parts.append("（空表）")
            if truncated:
                sheet_parts.append(
                    f"（该工作表超过 {_SHEET_ROW_LIMIT} 行，已截断，仅显示前 {_SHEET_ROW_LIMIT} 行）"
                )
            parts.append("\n".join(sheet_parts))
    finally:
        wb.close()
    text = "\n\n".join(parts)
    parse_time = time.perf_counter() - t0

    if file_hash:
        set_ocr_cache(file_hash, text)
        print(f"[Excel 缓存写入] file_hash={file_hash[:12]}... (解析耗时: {parse_time:.4f}s)")

    return text, {
        "ocr_cache_hit": False,
        "ocr_time": round(parse_time, 4),
        "ocr_cache_query_time": round(query_time * 1000, 2),
    }


# ------------------------------------------------------------------
# 旧版 .doc/.xls：LibreOffice headless 转换为 OOXML 后再走上面的解析器
# ------------------------------------------------------------------

async def _convert_with_libreoffice(src_path: Path, target_format: str) -> Path:
    """用 LibreOffice headless 将 .doc/.xls 转为 .docx/.xlsx，返回转换后文件路径。

    每次调用使用独立的 UserInstallation 目录，规避 soffice 并发 profile 锁冲突。
    缺失 soffice 或转换失败时抛 RuntimeError（由 /doc_text 转成 400 提示）。
    """
    if not LO_AVAILABLE:
        raise RuntimeError("LibreOffice(soffice) 未安装，无法解析旧版 .doc/.xls，请另存为 .docx/.xlsx 后上传")

    out_dir = Path(tempfile.mkdtemp(prefix="lo-conv-"))
    profile = f"file:///tmp/lo-profile-{uuid.uuid4().hex}"
    cmd = [
        "soffice", "--headless", "--nologo", "--nofirststartwizard",
        f"-env:UserInstallation={profile}",
        "--convert-to", target_format,
        "--outdir", str(out_dir),
        str(src_path),
    ]
    try:
        proc = await asyncio.create_subprocess_exec(
            *cmd,
            stdout=asyncio.subprocess.PIPE,
            stderr=asyncio.subprocess.PIPE,
        )
        stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=120)
        if proc.returncode != 0:
            raise RuntimeError(
                f"LibreOffice 转换失败（exit={proc.returncode}）: "
                f"{stderr.decode(errors='replace')[:300]}"
            )
        converted = out_dir / f"{src_path.stem}.{target_format}"
        if not converted.exists():
            raise RuntimeError(f"LibreOffice 未生成预期文件: {converted.name}")
        return converted
    except FileNotFoundError:
        raise RuntimeError("LibreOffice(soffice) 未安装，无法解析旧版 .doc/.xls，请另存为 .docx/.xlsx 后上传")
    except asyncio.TimeoutError:
        raise RuntimeError("LibreOffice 转换超时（>120s），请重试或另存为 .docx/.xlsx")


async def extract_doc_text(file_path: str, file_hash: str | None = None) -> tuple:
    """读取旧版 .doc：LibreOffice 转 .docx 后解析，按原 file_hash 缓存（命中免转换）。"""
    query_time = 0.0
    if file_hash:
        t0 = time.perf_counter()
        cached = get_ocr_cache(file_hash)
        query_time = time.perf_counter() - t0
        if cached is not None:
            return cached, {"ocr_cache_hit": True, "ocr_time": 0.0,
                            "ocr_cache_query_time": round(query_time * 1000, 2)}

    converted = await _convert_with_libreoffice(Path(file_path), "docx")
    try:
        text, _ = extract_docx_text(str(converted), None)  # None：不按临时文件 hash 缓存
    finally:
        try:
            converted.unlink(missing_ok=True)
            converted.parent.rmdir(ignore_errors=True)  # 临时目录为空时清理
        except OSError:
            pass

    if file_hash:
        set_ocr_cache(file_hash, text)  # 按原 .doc 的 hash 缓存，下次命中免转换
    return text, {
        "ocr_cache_hit": False,
        "ocr_time": 0.0,
        "ocr_cache_query_time": round(query_time * 1000, 2),
    }


async def extract_xls_text(file_path: str, file_hash: str | None = None) -> tuple:
    """读取旧版 .xls：LibreOffice 转 .xlsx 后解析，按原 file_hash 缓存（命中免转换）。"""
    query_time = 0.0
    if file_hash:
        t0 = time.perf_counter()
        cached = get_ocr_cache(file_hash)
        query_time = time.perf_counter() - t0
        if cached is not None:
            return cached, {"ocr_cache_hit": True, "ocr_time": 0.0,
                            "ocr_cache_query_time": round(query_time * 1000, 2)}

    converted = await _convert_with_libreoffice(Path(file_path), "xlsx")
    try:
        text, _ = extract_xlsx_text(str(converted), None)
    finally:
        try:
            converted.unlink(missing_ok=True)
            converted.parent.rmdir(ignore_errors=True)
        except OSError:
            pass

    if file_hash:
        set_ocr_cache(file_hash, text)
    return text, {
        "ocr_cache_hit": False,
        "ocr_time": 0.0,
        "ocr_cache_query_time": round(query_time * 1000, 2),
    }


# ------------------------------------------------------------------
# 字段提取取文本分发 + Excel 生成/模板填充助手
# ------------------------------------------------------------------

def _extract_text_for_concepts(file_path: str, file_hash: str | None) -> tuple:
    """按扩展名取目标文档全文，供 extract_concepts 喂给 LLM。

    PDF/图片 走 OCR；Word(.docx)/Excel(.xlsx) 走结构化解析（均同步，与原 PDF 路径同构）。
    旧版 .doc/.xls（异步 LibreOffice）暂不支持字段抽取，提示另存为 .docx/.xlsx。
    """
    ext = Path(file_path).suffix.lower()
    if ext == ".pdf":
        return extract_pdf_text(file_path, file_hash)
    if ext == ".docx":
        return extract_docx_text(file_path, file_hash)
    if ext == ".xlsx":
        return extract_xlsx_text(file_path, file_hash)
    if ext in _IMAGE_EXTS:
        return extract_image_text(file_path, file_hash)
    raise HTTPException(
        status_code=400,
        detail=f"字段提取暂不支持 {ext}，请上传 PDF/Word(.docx)/Excel(.xlsx)/图片(jpg/png/bmp/webp)",
    )


_XLSX_MIME = "application/vnd.openxmlformats-officedocument.spreadsheetml.sheet"


def _scan_template_fields(template_path: Path) -> tuple[dict, "Workbook"]:
    """扫描 Excel 模板，返回 ({field_name: [(ws, row, col), ...]}, wb)。

    规则（灵活布局）：非空字符串单元格，且其右侧单元格为空（且不在最末列）→ 视为
    「字段名 + 右侧空值位」。字段名做 strip + 去尾部冒号（：/:）归一化。
    天然过滤标题单元格（其右侧通常非空）。单次扫描，避免重复 I/O。
    """
    wb = load_workbook(template_path)  # read_write，保留公式与格式
    cells: dict[str, list] = {}
    for ws in wb.worksheets:
        max_col = ws.max_column
        for row in ws.iter_rows():
            for cell in row:
                if not isinstance(cell.value, str) or not cell.value.strip():
                    continue
                if cell.column >= max_col:  # 右侧超出已用范围 → 跳过（避免末列表头误判）
                    continue
                right = ws.cell(row=cell.row, column=cell.column + 1)
                if right.value in (None, ""):
                    name = cell.value.strip().rstrip("：:").strip()
                    if name:
                        cells.setdefault(name, []).append((ws, cell.row, cell.column + 1))
    return cells, wb


def _safe_set_cell(ws, row: int, col: int, value) -> None:
    """写入单元格；若目标位于合并区且非锚点则跳过，避免 openpyxl 写合并非锚点异常。"""
    for rng in ws.merged_cells.ranges:
        if rng.min_row <= row <= rng.max_row and rng.min_col <= col <= rng.max_col:
            if not (row == rng.min_row and col == rng.min_col):
                return  # 合并区非锚点，跳过该字段出现
    ws.cell(row=row, column=col, value=value)


def _fill_template(wb, cells: dict, results: dict) -> None:
    """把 results({field:value}) 写回模板各字段对应的右侧空值位。"""
    for name, targets in cells.items():
        v = results.get(name)
        val = str(v) if v is not None else ""
        for ws, r, c in targets:
            _safe_set_cell(ws, r, c, val)


def _build_default_excel(results: dict) -> Path:
    """{field:value} → 2 列「字段 | 值」默认 Excel，存 OUTPUT_DIR，返回路径。"""
    wb = Workbook()
    ws = wb.active
    ws.title = "提取结果"
    ws.append(["字段", "值"])
    ws.column_dimensions["A"].width = 24
    ws.column_dimensions["B"].width = 48
    for k, v in results.items():
        ws.append([str(k), str(v) if v is not None else ""])
    out = OUTPUT_DIR / generate_filename("extract_result.xlsx")
    wb.save(out)
    return out


def extract_concepts(file_path: str, fields: list, enhance: bool,
                     fields_enhance: list, fields_template: dict,
                     file_hash: str = None) -> tuple:
    """
    提取文档字段，支持 OCR 层缓存。
    file_hash: 由调用方预先计算并传入，避免重复计算。
    返回: (result_dict, timing_info)
    """

    # ---------------- 取全文（按扩展名分发，复用缓存） ----------------
    pdf_text, ocr_timing = _extract_text_for_concepts(file_path, file_hash)

    # 处理后的pdf内容
    with open(DATA_DIR / "temp.md", "w", encoding="UTF-8") as f:
        f.write(pdf_text)

    doc = Document(raw_text=pdf_text)
    doc.sat_model_id = SAT_MODEL_ID
    doc.paragraph_segmentation_mode = "sat"

    structure_dict = {field: str for field in fields}
    doc_concept = JsonObjectConcept(
        name="自定义概念抽取",
        description="抽取所有与指定概念相关的原文内容",
        structure=structure_dict,
        add_references=True,
        add_justifications=True,
        singular_occurrence=True,
    )
    doc.add_concepts([doc_concept])

    llm = DocumentLLM(
        model=OLLAMA_MODEL,
        api_base=OLLAMA_API_BASE,
        output_language="adapt",
        timeout=LLM_TIMEOUT,
        seed=LLM_SEED,
    )

    extracted_concepts = llm.extract_concepts_from_document(doc)

    res = {}
    if len(extracted_concepts[0].extracted_items) > 0:
        res = extracted_concepts[0].extracted_items[0].value

    if not res:
        print("+++++++++ 切换 concept 类型 ++++++++++")
        doc.concepts = [
            StringConcept(name=f, description=f"提取出文档中所有的{f}")
            for f in fields
        ]
        extracted_concepts = llm.extract_concepts_from_document(doc, max_items_per_call=50)
        for concept in extracted_concepts:
            res[concept.name] = ", ".join(item.value for item in concept.extracted_items)

    prompt = llm.get_usage()[-1].usage.calls[-1].prompt
    # 输入模型的文本提示
    with open(DATA_DIR / "prompt.txt", "w", encoding="UTF-8") as f:
        f.write(prompt)

    response = llm.get_usage()[-1].usage.calls[-1].response
    print("LLM 输出:")
    print(response)

    # 增强逻辑
    if enhance:
        print("++++++++++ 开始使用字段示例样本进行增强抽取 ++++++++++")
        examples = []
        if all(key in fields_template for key in fields_enhance):
            split_values = {
                key: [item.strip() for item in re.split(r"[;；]", fields_template[key]) if item.strip()]
                for key in fields_enhance
            }
            max_length = max(len(vals) for vals in split_values.values())
            padded_values = {}
            for key, vals in split_values.items():
                padded = vals + [vals[-1] if vals else ""] * (max_length - len(vals))
                padded_values[key] = padded
            for i in range(max_length):
                content = {key: padded_values[key][i] for key in fields_enhance}
                examples.append(JsonObjectExample(content=content))
        else:
            print("字段增强失败: fields_enhance 中的字段未全部存在于 fields_template 中。")

        res_enhance = extract_concepts_enhance(
            file=pdf_text,
            concepts=fields_enhance,
            concepts_template=examples,
        )
        for item_enhance in res_enhance:
            for k, v in item_enhance.value.items():
                if k in fields_enhance:
                    res[k] = v

    print("===============================")
    print(res)

    # OCR 计时由 extract_pdf_text 统一返回
    return res, ocr_timing


# ------------------------------------------------------------------
# 执行文档提取操作
# ------------------------------------------------------------------
def do_extract(request: DocEtractRequest, upload_dir: Path) -> dict:
    file_path = _resolve_upload(request.filename, upload_dir)
    if not file_path.exists():
        raise FileNotFoundError("文件不存在")

    # 计算缓存键
    hash_start = time.perf_counter()
    file_hash = compute_file_hash(str(file_path))
    fields_hash = compute_fields_hash(
        request.fields, request.enhance, request.fields_enhance, request.fields_template
    )
    hash_time = time.perf_counter() - hash_start

    # ---------------- LLM 层缓存查询 ----------------
    llm_query_start = time.perf_counter()
    cached_result = get_llm_cache(file_hash, fields_hash)
    llm_query_time = time.perf_counter() - llm_query_start

    if cached_result is not None:
        total_time = time.perf_counter() - hash_start
        print(f"[LLM 缓存命中] file_hash={file_hash[:12]}..., fields_hash={fields_hash[:12]}...")
        print(f"[耗时对比] 缓存命中总耗时: {total_time*1000:.2f}ms (哈希计算: {hash_time*1000:.2f}ms, 缓存查询: {llm_query_time*1000:.2f}ms)")
        return {
            "message": "提取成功（LLM缓存命中）",
            "results": cached_result,
            "timing": {
                "cache_hit": "llm",
                "total_time_ms": round(total_time * 1000, 2),
                "hash_time_ms": round(hash_time * 1000, 2),
                "llm_cache_query_ms": round(llm_query_time * 1000, 2),
            },
        }

    print()
    print("++++++++++ 开始进行文档提取 ++++++++++")
    start = time.perf_counter()
    res, extract_timing = extract_concepts(
        str(file_path), request.fields, request.enhance,
        request.fields_enhance, request.fields_template,
        file_hash=file_hash,
    )
    end = time.perf_counter()
    total_time = end - start
    print(f"提取用时: {total_time:.4f} 秒")

    # ---------------- 写入 LLM 缓存 ----------------
    set_llm_cache(file_hash, fields_hash, res)
    print(f"[LLM 缓存写入] file_hash={file_hash[:12]}..., fields_hash={fields_hash[:12]}...")

    ocr_hit_str = "是" if extract_timing["ocr_cache_hit"] else "否"
    print(f"[耗时对比] 本次总耗时: {total_time*1000:.2f}ms | OCR缓存命中: {ocr_hit_str} | OCR耗时: {extract_timing['ocr_time']*1000:.2f}ms | LLM+其他: {(total_time - extract_timing['ocr_time'])*1000:.2f}ms")

    return {
        "message": "提取成功",
        "results": res,
        "timing": {
            "cache_hit": "ocr" if extract_timing["ocr_cache_hit"] else "none",
            "total_time_ms": round(total_time * 1000, 2),
            "hash_time_ms": round(hash_time * 1000, 2),
            "ocr_cache_hit": extract_timing["ocr_cache_hit"],
            "ocr_time_ms": round(extract_timing["ocr_time"] * 1000, 2),
            "ocr_cache_query_ms": extract_timing["ocr_cache_query_time"],
        },
    }


@app.post("/doc_text", summary="获取上传文档的全文文本")
async def doc_text(request: DocTextRequest,
                   current_user: str = Depends(verify_token)) -> Dict[str, Any]:
    """返回已上传文档的全文，供 agent 做问答/摘要/解读等通用处理。

    按扩展名分发：PDF 走 OCR；Word(.docx)/Excel(.xlsx) 走结构化解析；
    旧版 .doc/.xls 经 LibreOffice 转 OOXML 后再解析。各类全文共用 ocr_cache
    文本缓存（按 file_hash 键），已读取过的文件秒回。
    """
    file_path = _resolve_upload(request.filename)
    if not file_path.exists():
        raise HTTPException(status_code=400, detail="文件不存在")

    suffix = file_path.suffix.lower()
    try:
        file_hash = compute_file_hash(str(file_path))
        if suffix == ".pdf":
            text, timing = extract_pdf_text(str(file_path), file_hash)
        elif suffix == ".docx":
            text, timing = extract_docx_text(str(file_path), file_hash)
        elif suffix == ".xlsx":
            text, timing = extract_xlsx_text(str(file_path), file_hash)
        elif suffix == ".doc":
            text, timing = await extract_doc_text(str(file_path), file_hash)
        elif suffix == ".xls":
            text, timing = await extract_xls_text(str(file_path), file_hash)
        elif suffix in _IMAGE_EXTS:
            text, timing = extract_image_text(str(file_path), file_hash)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"不支持的文件类型: {suffix}，支持 PDF/Word(.docx/.doc)/Excel(.xlsx/.xls)/图片(jpg/png/bmp/webp)",
            )
    except HTTPException:
        raise
    except RuntimeError as e:
        # LibreOffice 未安装/转换失败等：给出可读提示
        traceback.print_exc()
        raise HTTPException(status_code=400, detail=str(e))
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"读取文档文本失败：{e}")

    return {
        "message": "读取成功",
        "filename": request.filename,
        "text": text,
        "chars": len(text),
        "cache_hit": timing["ocr_cache_hit"],
    }


@app.post("/image_text", summary="OCR 识别图片文本")
async def image_text(file: UploadFile = File(...),
                     current_user: str = Depends(verify_token)) -> Dict[str, Any]:
    """对上传图片做 OCR，返回识别出的文本（Markdown）。复用 PPStructure + file_hash 缓存。

    图片由调用方（office_mcp read_image）以 multipart 上传，写到临时文件后 OCR，
    不落入 UPLOAD_DIR（图片存储在 agent 侧），结束后清理临时文件。
    """
    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in _IMAGE_EXTS:
        raise HTTPException(status_code=400, detail="仅支持图片文件（jpg/jpeg/png/bmp/webp）")

    tmp = tempfile.NamedTemporaryFile(suffix=suffix, delete=False)
    tmp_path = tmp.name
    try:
        while chunk := await file.read(1024 * 1024):
            tmp.write(chunk)
        tmp.close()  # 落盘并释放写句柄，供 compute_file_hash 与 PPStructure 读取
        file_hash = compute_file_hash(tmp_path)
        text, timing = extract_image_text(tmp_path, file_hash)
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"图片识别失败：{e}")
    finally:
        try:
            tmp.close()  # 幂等：写入阶段异常时确保关闭
        except Exception:
            pass
        try:
            os.remove(tmp_path)
        except OSError:
            pass

    return {
        "message": "识别成功",
        "filename": file.filename,
        "text": text,
        "chars": len(text),
        "cache_hit": timing["ocr_cache_hit"],
    }


@app.post("/doc_extract", summary="执行文档提取操作")
async def doc_extract(request: DocEtractRequest,
                      current_user: str = Depends(verify_token)) -> Dict[str, Any]:
    if request.fields is None:
        raise HTTPException(status_code=400, detail="待抽取的字段列表fields不能为空")

    if len(request.fields) != len(set(request.fields)):
        raise HTTPException(status_code=400, detail="待抽取的字段列表fields中存在重复字段")

    if request.enhance:
        if request.fields_enhance is None:
            raise HTTPException(status_code=400, detail="增强字段列表fields_enhance不能为空")
        if any(v is None or (isinstance(v, (list, tuple, dict, str)) and len(v) == 0) for v in request.fields_template.values()):
            raise HTTPException(status_code=400, detail="使用增强抽取，字段示例样本fields_template不能为空")

        extra = set(request.fields_enhance) - set(request.fields)
        if extra:
            raise HTTPException(status_code=400, detail="增强字段列表fields_enhance必须是字段列表field的子集!")

        # fields_enhance 与 fields_template 的 key 必须严格一一对应
        if set(request.fields_enhance) != set(request.fields_template.keys()):
            raise HTTPException(status_code=400, detail="增强字段列表的元素和字段示例样本的key必须一一对应!")

    try:
        # 直接同步调用，避免 ThreadPoolExecutor 导致 CUDA 上下文跨线程报错
        result = do_extract(request, UPLOAD_DIR)
    except HTTPException:
        # 透传 do_extract 内 _resolve_upload 的 400（非法文件名），不被下方 500 吞掉
        raise
    except FileNotFoundError as e:
        raise HTTPException(status_code=400, detail=f"抽取失败：{str(e)}")
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"抽取失败：{e}")

    return result


# ------------------------------------------------------------------
# 字段提取 → Excel 下载
# ------------------------------------------------------------------

class ExtractExcelRequest(BaseModel):
    filename: str
    fields: List[str]
    enhance: bool = False
    fields_enhance: List[str] = []
    fields_template: dict = {}
    display_name: str | None = None  # 下载文件名（可选）


class FillTemplateRequest(BaseModel):
    filename: str            # 目标文档 saved_name
    template_filename: str   # Excel 模板 saved_name
    enhance: bool = False
    fields_enhance: List[str] = []
    fields_template: dict = {}
    display_name: str | None = None


@app.get("/download/{filename}", summary="下载生成的 Excel")
async def download_excel(filename: str, name: str | None = None):
    """公开下载 OUTPUT_DIR 下的生成 Excel（文件名为不可猜 uuid）。

    供前端页面与 agent（经 office_mcp 返回的 download_url）直接浏览器下载；
    不要求 JWT（与 document_compare 的 /get_file 一致），靠不可猜文件名保护。
    """
    path = _resolve_upload(filename, OUTPUT_DIR)  # 复用路径穿越防御，base=OUTPUT_DIR
    if not path.exists():
        raise HTTPException(status_code=404, detail="文件不存在")
    disp = name or filename
    return FileResponse(
        path=path,
        media_type=_XLSX_MIME,
        headers={"Content-Disposition": f'attachment; filename="{quote(disp)}"'},
    )


@app.post("/extract_to_excel", summary="提取字段→默认 Excel 下载")
async def extract_to_excel(req: ExtractExcelRequest,
                           current_user: str = Depends(verify_token)) -> Dict[str, Any]:
    """从目标文档提取指定字段，生成 2 列「字段|值」默认 Excel 供下载。

    复用 do_extract（含 LLM/OCR 缓存）。目标文档支持 PDF/Word(.docx)/Excel(.xlsx)。
    """
    fields = [f.strip() for f in req.fields if f and f.strip()]
    if not fields:
        raise HTTPException(status_code=400, detail="字段列表不能为空")
    dr = DocEtractRequest(
        filename=req.filename, fields=fields, enhance=req.enhance,
        fields_enhance=req.fields_enhance, fields_template=req.fields_template,
    )
    try:
        res = do_extract(dr, UPLOAD_DIR)
    except HTTPException:
        raise
    except FileNotFoundError as e:
        raise HTTPException(status_code=400, detail=f"提取失败：{str(e)}")
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"提取失败：{e}")

    out = _build_default_excel(res["results"])
    url = f"{PUBLIC_BASE_URL}/download/{out.name}?name={quote(req.display_name or '字段提取结果.xlsx')}"
    return {
        "message": "提取成功",
        "filename": out.name,
        "download_url": url,
        "results": res["results"],
        "cache_hit": res["timing"].get("cache_hit"),
    }


@app.post("/fill_template", summary="按 Excel 模板提取并填充→下载")
async def fill_template(req: FillTemplateRequest,
                        current_user: str = Depends(verify_token)) -> Dict[str, Any]:
    """按用户上传的 Excel 模板「字段名 + 右侧空单元格」识别字段，从目标文档提取后
    把值填回各字段右侧位置，另存为 Excel 供下载（不覆盖原模板）。

    目标文档支持 PDF/Word(.docx)/Excel(.xlsx)；模板仅支持 .xlsx（.xls 需另存为 .xlsx）。
    """
    template_path = _resolve_upload(req.template_filename, UPLOAD_DIR)
    if template_path.suffix.lower() not in (".xlsx", ".xls"):
        raise HTTPException(status_code=400, detail="模板须为 Excel(.xlsx/.xls)")
    if template_path.suffix.lower() == ".xls":
        raise HTTPException(
            status_code=400,
            detail="旧版 .xls 模板暂不支持，请另存为 .xlsx 后上传",
        )
    if not template_path.exists():
        raise HTTPException(status_code=400, detail="模板文件不存在")

    try:
        cells, wb = _scan_template_fields(template_path)
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=400, detail=f"模板解析失败：{e}")
    if not cells:
        raise HTTPException(
            status_code=400,
            detail="模板未识别到待填字段（需含「字段名 + 右侧空单元格」结构）",
        )

    dr = DocEtractRequest(
        filename=req.filename, fields=list(cells.keys()), enhance=req.enhance,
        fields_enhance=req.fields_enhance, fields_template=req.fields_template,
    )
    try:
        res = do_extract(dr, UPLOAD_DIR)
    except HTTPException:
        raise
    except FileNotFoundError as e:
        raise HTTPException(status_code=400, detail=f"提取失败：{str(e)}")
    except Exception as e:
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"提取失败：{e}")

    _fill_template(wb, cells, res["results"])
    out = OUTPUT_DIR / generate_filename("filled.xlsx")
    wb.save(out)
    url = f"{PUBLIC_BASE_URL}/download/{out.name}?name={quote(req.display_name or '模板填充结果.xlsx')}"
    return {
        "message": "填充成功",
        "filename": out.name,
        "download_url": url,
        "results": res["results"],
        "fields_found": list(cells.keys()),
        "cache_hit": res["timing"].get("cache_hit"),
    }


@app.post("/doc_test", summary="测试接口")
async def doc_test(req: DocEtractRequest):
    print(req)
    return {"message": "成功", "status_code": 200, "data": req}


if __name__ == "__main__":
    uvicorn.run("main:app", host="0.0.0.0", port=PORT, reload=False)
